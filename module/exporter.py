from io import BytesIO
import math

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

def generate_upsc_docx(questions):
    if Document is None:
        return None

    doc = Document()
    
    # Title
    heading = doc.add_heading('UPSC Prelims Mock Test', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Time Allowed: 2 Hours\t\t\t\tMaximum Marks: ---')
    doc.add_paragraph('_'*80)
    doc.add_paragraph()
    
    # Questions
    for i, q in enumerate(questions):
        # Hindi Question
        p_h = doc.add_paragraph()
        p_h.add_run(f"{i+1}. ").bold = True
        p_h.add_run(q.question_hindi)
        
        # Hindi Options
        for idx, opt in enumerate(q.options_hindi):
             doc.add_paragraph(f"({chr(97+idx)}) {opt}")
        
        doc.add_paragraph() # Gap
        
        # English Question
        p_e = doc.add_paragraph()
        p_e.add_run(f"{i+1}. ").bold = True
        p_e.add_run(q.question_english)
        
        # English Options
        for idx, opt in enumerate(q.options_english):
             doc.add_paragraph(f"({chr(97+idx)}) {opt}")
        
        # Add Answer after each question
        doc.add_paragraph(f"Answer: {q.answer}")
             
        doc.add_paragraph('_'*40) # Separator
        doc.add_paragraph() # Spacer

    # Answer Key (New Page)
    doc.add_page_break()
    doc.add_heading('Answer Key', level=1)
    
    # Configure Table for Vertical Layout with Fixed 20 Rows Limit per Column
    LIMIT_PER_COL = 20
    
    total_questions = len(questions)
    if total_questions == 0:
        # Save empty
        f = BytesIO()
        doc.save(f)
        f.seek(0)
        return f

    num_pairs = math.ceil(total_questions / LIMIT_PER_COL)
    
    # Table height is 20 rows max, or less if total questions < 20
    table_rows_count = min(total_questions, LIMIT_PER_COL)
    
    total_cols = num_pairs * 2
    
    table = doc.add_table(rows=table_rows_count + 1, cols=total_cols) # +1 for Header
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    for k in range(num_pairs):
        hdr_cells[2*k].text = 'Q.No'
        hdr_cells[2*k+1].text = 'Ans'
        
        # Make header bold
        for cell_idx in [2*k, 2*k+1]:
            run = hdr_cells[cell_idx].paragraphs[0].runs
            if run: run[0].bold = True
            else: hdr_cells[cell_idx].paragraphs[0].add_run(hdr_cells[cell_idx].text).bold = True

    # Fill Data
    for q_idx in range(total_questions):
        # Determine which column pair (k) and which row (r)
        k = q_idx // LIMIT_PER_COL
        r = q_idx % LIMIT_PER_COL
        
        q = questions[q_idx]
        
        # Table row index is r + 1 (because of header)
        if r + 1 < len(table.rows):
            row_cells = table.rows[r + 1].cells
            
            # Ensure we don't access out of bounds columns (though logic ensures total_cols covers it)
            if 2*k+1 < len(row_cells):
                row_cells[2*k].text = str(q_idx + 1)
                row_cells[2*k+1].text = q.answer

    # Save to BytesIO
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f
